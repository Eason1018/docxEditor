import os
from docx import Document
import csv


def analyze_document(docx_file):
    """
    Analyze the structure of the Word document (e.g., tables) and log it.
    """
    print(f"Analyzing document: {docx_file}")
    doc = Document(docx_file)

    # Analyze all tables in the document
    for table_index, table in enumerate(doc.tables):
        print(f"\nTable {table_index}:")
        for row_index, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f" Row {row_index}: {row_data}")


def populate_table_0(table):
    """
    Populate specific fields in Table 0 with hardcoded data, handling shared rows.
    """
    # Hardcoded data for Table 0
    table_0_data = {
        "From:": "Ng, Wai Ming, Rock / 吳偉明 / 88888",
        "To:": "Tendering Committee",
        "Name of Property:": "AP- Apec Plaza",
        "The Works:": "電機控制系統問題扶手鬆動滑輪/滑道磨損安全營示系統失靈：運行並確認問題維修過程需要專業技術人員進行,確保電梯/扶手電梯安全可靠運行",
        "Tender Ref.:": "SHKP1234-001"
    }

    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            for i, cell in enumerate(cells):
                if cell.text.strip() in table_0_data:
                    key = cell.text.strip()
                    cells[i + 1].text = table_0_data[key]  # Place value in the next cell


def populate_table_from_csv(doc, csv_file, table_index=1):
    """
    Populate Table 1 using data from a CSV file, handling empty rows and column alignment.
    """
    table = doc.tables[table_index]

    # Map CSV headers to table column indices
    header_map = {
        "Serial No.": 0,
        "Tenderer Name": 1,
        "Tenderer Notified On": 2,
    }

    # Read CSV data
    with open(csv_file, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        headers = next(reader)  # Extract CSV header row

        # Populate rows in the table
        row_idx = 2  # Start after the header rows
        for row_data in reader:
            while row_idx < len(table.rows):
                table_row = table.rows[row_idx]
                if all(not cell.text.strip() for cell in table_row.cells):  # Skip empty rows
                    row_idx += 1
                    continue

                for csv_col, table_col in header_map.items():
                    if table_col < len(table_row.cells):
                        table_row.cells[table_col].text = row_data[headers.index(csv_col)].strip()
                row_idx += 1
                break


# File paths
INPUT_DOCX = "input.docx"  # Replace with your actual .docx file path
OUTPUT_DOCX = "output.docx"  # Path for the modified .docx file
INPUT_CSV = "data.csv"  # Path to the CSV file

# Ensure file exists
if not os.path.exists(INPUT_DOCX):
    print(f"Error: Input file '{INPUT_DOCX}' does not exist.")
    exit(1)

if not os.path.exists(INPUT_CSV):
    print(f"Error: Input file '{INPUT_CSV}' does not exist.")
    exit(1)

# Delete the existing output file if it exists
if os.path.exists(OUTPUT_DOCX):
    print(f"Deleting existing output file: {OUTPUT_DOCX}")
    os.remove(OUTPUT_DOCX)

# Step 1: Analyze the input document structure
print("\nSTEP 1: Analyzing Input Document")
analyze_document(INPUT_DOCX)

# Load the Word document
doc = Document(INPUT_DOCX)

# Step 2: Populate Table 0 with hardcoded data
print("\nSTEP 2: Populating Table 0 with Hardcoded Data...")
try:
    populate_table_0(doc.tables[0])
    print("Table 0 populated successfully.")
except Exception as e:
    print(f"Error during Table 0 population: {e}")

# Step 3: Populate Table 1 using CSV data
print("\nSTEP 3: Populating Table 1 from CSV...")
try:
    populate_table_from_csv(doc, INPUT_CSV, table_index=1)
    print("Table 1 populated successfully.")
except Exception as e:
    print(f"Error during Table 1 population: {e}")

# Save the modified document
doc.save(OUTPUT_DOCX)

# Step 4: Analyze the output document to verify changes
print("\nSTEP 4: Analyzing Output Document")
analyze_document(OUTPUT_DOCX)

print(f"\nProcess completed. Modified document saved as: {OUTPUT_DOCX}")
