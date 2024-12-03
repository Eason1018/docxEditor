import os
import zipfile
import csv
from docx.shared import Inches
from PIL import Image
from PIL import ImageOps
from docx.oxml import OxmlElement
from lxml import etree
from docx import Document
import pypandoc
import win32com.client

def extract_docx(docx_path, extract_dir):
    """
    Extracts a .docx file into a specified directory.
    """
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)


def analyze_document(docx_file):
    """
    Analyze the structure of the Word document (e.g., tables) and log it,
    including handling merged cells and skipping empty rows.
    """
    print(f"Analyzing document: {docx_file}")
    doc = Document(docx_file)

    for table_index, table in enumerate(doc.tables):
        print(f"\nTable {table_index}:")
        for row_index, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # Check if the cell is merged or empty
                cell_text = cell.text.strip()
                if not cell_text:
                    cell_text = "<empty>"

                # Append cell data to the row
                row_data.append(cell_text)

            # Skip rows that are completely empty
            if all(cell == "<empty>" for cell in row_data):
                print(f" Row {row_index}: <empty row>")
            else:
                print(f" Row {row_index}: {row_data}")


def add_signature_to_cell(cell, image_path):
    """
    Add a signature image to a table cell, resize it, and adjust the row height.
    :param cell: The table cell where the image will be added.
    :param image_path: Path to the signature image file.
    """
    try:
        print(f"Adding signature to cell: {image_path}")
        # Clear existing text
        cell.text = ""

        # Get the approximate cell width
        cell_width = get_cell_width(cell)
        if not cell_width:
            cell_width = 1.0  # Default to 1 inch if width cannot be determined

        print(f"Cell width: {cell_width} inches")

        # Open the image and resize it to fit the cell
        with Image.open(image_path) as img:
            aspect_ratio = img.width / img.height
            img_width = cell_width * 96  # Convert inches to pixels (assuming 96 DPI)
            img_height = img_width / aspect_ratio

            print(f"Image resized to width: {img_width}px, height: {img_height}px")

            # Save resized image to a temporary file
            resized_image_path = "resized_signature.png"
            img = img.resize((int(img_width), int(img_height)), Image.Resampling.LANCZOS)
            img.save(resized_image_path)

        # Add the resized image to the cell
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(resized_image_path, width=Inches(cell_width))

        # Adjust the row height to fit the image
        adjust_row_height(cell._tc, img_height)
        print(f"Signature added successfully to cell.")
    except Exception as e:
        print(f"Error adding signature: {e}")

def adjust_row_height(tc, img_height_px):
    """
    Adjust the row height to fit the signature and disable auto height adjustment.
    :param tc: The table cell's underlying XML element.
    :param img_height_px: The image height in pixels.
    """
    try:
        print(f"Adjusting row height for image height: {img_height_px}px")
        img_height_twips = int(img_height_px * 15)  # Convert pixels to twips
        print(f"Calculated row height: {img_height_twips} twips")

        # Get the parent row
        tr = tc.getparent()
        trPr = tr.get_or_add_trPr()

        # Set row height
        rowHeight = OxmlElement('w:trHeight')
        rowHeight.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(img_height_twips))
        rowHeight.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hRule', 'exact')
        trPr.append(rowHeight)

        print(f"Row height adjusted to: {img_height_twips} twips")
    except Exception as e:
        print(f"Error adjusting row height: {e}")


def get_cell_width(cell):
    """
    Estimate the width of the cell in inches.
    """
    try:
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find("w:tcW", namespaces)  # Use namespace
        if tcW is not None:
            width_twips = int(tcW.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"))
            width_inches = width_twips / 1440  # Convert twips to inches
            return width_inches
        else:
            print("Cell width could not be determined; using default.")
            return None
    except Exception as e:
        print(f"Error calculating cell width: {e}")
        return None



def modify_document_xml(document_xml_path, replacements):
    """
    Modifies the document.xml file by replacing specified text.
    """
    parser = etree.XMLParser(ns_clean=True, recover=True)
    tree = etree.parse(document_xml_path, parser)
    root = tree.getroot()
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for text_element in root.xpath('.//w:t', namespaces=namespaces):
        if text_element.text and any(old_text in text_element.text for old_text in replacements):
            for old_text, new_text in replacements.items():
                if old_text in text_element.text:
                    text_element.text = text_element.text.replace(old_text, new_text)

    tree.write(document_xml_path, xml_declaration=True, encoding='UTF-8', standalone="yes")


def repack_docx(extract_dir, output_docx_path):
    """
    Repackages the extracted files back into a .docx file.
    """
    with zipfile.ZipFile(output_docx_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root_dir, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path = os.path.join(root_dir, file)
                arcname = os.path.relpath(file_path, extract_dir)
                zipf.write(file_path, arcname)


def populate_table_0(table):
    """
    Populate specific fields in Table 0 with hardcoded data.
    """
    table_0_data = {
        "From:": "Ng, Wai Ming, Rock / 吳偉明 / 88888",
        "To:": "Tendering Committee",
        "Name of Property:": "AP- Apec Plaza",
        "The Works:": "電機控制系統問題扶手鬆動滑輪/滑道磨損安全營示系統失靈：運行並確認問題維修過程需要專業技術人員進行,確保電梯/扶手電梯安全可靠運行",
        "Tender Ref.:": "SHKP1234-001"
    }

    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            for i, cell in enumerate(cells):
                if cell.text.strip() in table_0_data:
                    key = cell.text.strip()
                    cells[i + 1].text = table_0_data[key]


def populate_table_from_csv(doc, csv_file, table_index=1):
    """
    Populate Table 1 using data from a CSV file.
    """
    table = doc.tables[table_index]

    with open(csv_file, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)  # Use DictReader for column matching

        csv_rows = list(reader)  # Load all rows into memory for better control
        csv_index = 0  # Track current CSV row being used

        for row_index, table_row in enumerate(table.rows):
            # Stop if all CSV rows are consumed
            if csv_index >= len(csv_rows):
                break

            # Check if this row has empty cells to populate
            if all(not cell.text.strip() for cell in table_row.cells):
                csv_data = csv_rows[csv_index]  # Get current CSV row

                # Populate cells based on CSV columns
                for col_name, cell in zip(reader.fieldnames, table_row.cells):
                    if col_name in csv_data:
                        cell.text = csv_data[col_name]

                csv_index += 1  # Move to the next CSV row



def add_row_to_table(doc, table_index, row_data):
    """
    Add a new row to the specified table with the given data and maintain formatting.
    """
    table = doc.tables[table_index]

    # Use the last row as a template for formatting
    template_row = table.rows[-1]

    # Add a new row
    new_row = table.add_row()

    # Populate the new row with the provided data
    for i, cell_data in enumerate(row_data):
        if i < len(new_row.cells):  # Ensure the data fits in the table
            new_cell = new_row.cells[i]
            new_cell.text = cell_data

            # Copy formatting from the template row
            template_cell = template_row.cells[i] if i < len(template_row.cells) else None
            if template_cell:
                _copy_cell_formatting(template_cell, new_cell)

    print(f"Row added to Table {table_index} with data: {row_data}")


def _copy_cell_formatting(template_cell, target_cell):
    """
    Copy the text formatting from the template cell to the target cell.
    """
    # Copy paragraph alignment
    if template_cell.paragraphs and target_cell.paragraphs:
        target_cell.paragraphs[0].alignment = template_cell.paragraphs[0].alignment
    template_paragraph = template_cell.paragraphs[0]
    target_paragraph = target_cell.paragraphs[0]

    # Copy font settings
    if template_paragraph.runs and target_paragraph.runs:
        template_font = template_paragraph.runs[0].font
        target_font = target_paragraph.runs[0].font

        # Copy font size, name, and other attributes
        if template_font.size:
            target_font.size = template_font.size
        if template_font.name:
            target_font.name = template_font.name
        target_font.bold = template_font.bold
        target_font.italic = template_font.italic
        target_font.underline = template_font.underline

    # Copy alignment (if needed)
    target_paragraph.alignment = template_paragraph.alignment



def delete_row_from_table(doc, table_index, row_number):
    """
    Delete a specific row from the specified table.
    """
    table = doc.tables[table_index]

    if row_number < 0 or row_number >= len(table.rows):
        print(f"Invalid row number: {row_number}")
        return

    # Access the row to delete
    row_to_delete = table.rows[row_number]

    # Remove the row from the table
    tbl = table._tbl
    tbl.remove(row_to_delete._tr)

    print(f"Row {row_number} deleted from Table {table_index}.")


def convert_to_pdf(input_docx, output_pdf):
    """
    Convert a .docx file to .pdf using win32com.client (requires Microsoft Word).
    """
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(input_docx)
        doc.SaveAs(output_pdf, FileFormat=17)  # 17 is the constant for wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"PDF file created: {output_pdf}")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
